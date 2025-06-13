#!/usr/bin/env python3
"""
ExcelからWordへ文章内容をコピーし、PDFで出力するプログラム
"""

import os
import sys
from pathlib import Path
from typing import List, Optional, Union
import argparse

# Excel操作用
from openpyxl import load_workbook
from openpyxl.cell.cell import Cell
from openpyxl.utils import get_column_letter

# Word操作用
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# PDF変換用
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.platypus.tables import Table
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY


class ExcelToWordPDFConverter:
    """ExcelファイルをWord経由でPDFに変換するクラス"""
    
    def __init__(self):
        self.styles = getSampleStyleSheet()
        self._setup_japanese_font()
    
    def _setup_japanese_font(self):
        """日本語フォントの設定（システムフォントを使用）"""
        try:
            # 日本語フォントのスタイルを作成
            self.japanese_style = ParagraphStyle(
                'Japanese',
                parent=self.styles['Normal'],
                fontName='Helvetica',
                fontSize=10,
                leading=12,
            )
        except Exception as e:
            print(f"フォント設定エラー: {e}")
            self.japanese_style = self.styles['Normal']
    
    def read_excel(self, excel_path: str) -> List[List[str]]:
        """Excelファイルからデータを読み取る"""
        try:
            workbook = load_workbook(excel_path, data_only=True)
            sheet = workbook.active
            
            data = []
            for row in sheet.iter_rows():
                row_data = []
                for cell in row:
                    value = cell.value if cell.value is not None else ""
                    row_data.append(str(value))
                if any(row_data):  # 空行でない場合のみ追加
                    data.append(row_data)
            
            workbook.close()
            return data
            
        except Exception as e:
            print(f"Excelファイルの読み取りエラー: {e}")
            raise
    
    def create_word_document(self, data: List[List[str]], word_path: str):
        """データからWordドキュメントを作成"""
        try:
            doc = Document()
            
            # タイトルを追加（最初の行をタイトルとして扱う）
            if data:
                doc.add_heading('Excel Data Export', 0)
                
                # テーブル形式でデータを追加
                if len(data) > 0:
                    # 最大列数を計算
                    max_cols = max(len(row) for row in data)
                    
                    # テーブルを作成
                    table = doc.add_table(rows=len(data), cols=max_cols)
                    table.style = 'Light Grid Accent 1'
                    
                    # データをテーブルに追加
                    for i, row_data in enumerate(data):
                        for j, cell_data in enumerate(row_data):
                            if j < max_cols:
                                table.cell(i, j).text = cell_data
                
                # 段落として追加する場合のコード（コメントアウト）
                # for row in data:
                #     p = doc.add_paragraph()
                #     p.add_run(' | '.join(row))
            
            doc.save(word_path)
            print(f"Wordドキュメントを作成しました: {word_path}")
            
        except Exception as e:
            print(f"Wordドキュメントの作成エラー: {e}")
            raise
    
    def convert_to_pdf_from_data(self, data: List[List[str]], pdf_path: str):
        """データから直接PDFを作成（Word経由せず）"""
        try:
            doc = SimpleDocTemplate(pdf_path, pagesize=A4)
            story = []
            
            # タイトル
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=self.styles['Heading1'],
                fontSize=16,
                textColor=colors.HexColor('#000080'),
                spaceAfter=30,
                alignment=TA_CENTER
            )
            story.append(Paragraph("Excel Data Export", title_style))
            story.append(Spacer(1, 12))
            
            if data:
                # テーブル形式でデータを追加
                table_data = []
                for row in data:
                    # 各セルをParagraphオブジェクトに変換（長いテキストの折り返し対応）
                    table_row = []
                    for cell in row:
                        p = Paragraph(cell, self.japanese_style)
                        table_row.append(p)
                    # 不足している列を空文字で埋める
                    max_cols = max(len(r) for r in data)
                    while len(table_row) < max_cols:
                        table_row.append(Paragraph("", self.japanese_style))
                    table_data.append(table_row)
                
                # テーブルを作成
                table = Table(table_data)
                
                # テーブルスタイルを設定
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 12),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ]))
                
                story.append(table)
            
            # PDFを生成
            doc.build(story)
            print(f"PDFファイルを作成しました: {pdf_path}")
            
        except Exception as e:
            print(f"PDF作成エラー: {e}")
            raise
    
    def convert(self, excel_path: str, output_dir: str = None):
        """ExcelファイルをWordとPDFに変換する"""
        # パスの設定
        excel_path = Path(excel_path)
        if not excel_path.exists():
            raise FileNotFoundError(f"Excelファイルが見つかりません: {excel_path}")
        
        if output_dir is None:
            output_dir = excel_path.parent
        else:
            output_dir = Path(output_dir)
            output_dir.mkdir(parents=True, exist_ok=True)
        
        # 出力ファイル名の設定
        base_name = excel_path.stem
        word_path = output_dir / f"{base_name}.docx"
        pdf_path = output_dir / f"{base_name}.pdf"
        
        # 処理の実行
        print(f"Excelファイルを処理中: {excel_path}")
        
        # Excelデータを読み取る
        data = self.read_excel(str(excel_path))
        
        # Wordドキュメントを作成
        self.create_word_document(data, str(word_path))
        
        # PDFに変換
        self.convert_to_pdf_from_data(data, str(pdf_path))
        
        return str(word_path), str(pdf_path)


def main():
    """メイン関数"""
    parser = argparse.ArgumentParser(description='ExcelファイルをWord経由でPDFに変換します')
    parser.add_argument('excel_file', help='変換するExcelファイル')
    parser.add_argument('-o', '--output', help='出力ディレクトリ（省略時は入力ファイルと同じディレクトリ）')
    
    args = parser.parse_args()
    
    try:
        converter = ExcelToWordPDFConverter()
        word_path, pdf_path = converter.convert(args.excel_file, args.output)
        
        print("\n変換完了!")
        print(f"Word: {word_path}")
        print(f"PDF: {pdf_path}")
        
    except Exception as e:
        print(f"\nエラーが発生しました: {e}")
        sys.exit(1)


if __name__ == "__main__":
    main()